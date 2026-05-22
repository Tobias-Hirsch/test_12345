# Developer: Jinglu Han
# mailbox: admin@de-manufacturing
import aiohttp
import io
import os
import tempfile
from docx import Document
from typing import Optional
import traceback
from app.services.logging import word_process_content_logger as logger

# Kommentar
try:
    import docx2txt
    DOCX2TXT_AVAILABLE = True
except ImportError:
    DOCX2TXT_AVAILABLE = False
    logger.warning("docx2txtWarnhinweis")


async def extract_word_content_from_url(url: str) -> Optional[str]:
    """
    Hinweis

    Args:
        url: WordDokumenteHinweis

    Returns:
        str: Hinweis
    """
    temp_file_path = None
    try:
        # Kommentar
        headers = {
            'User-Agent': 'Mozilla/5.0 (Macintosh; Intel Mac OS X 10_15_7) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/91.0.4472.124 Safari/537.36',
            'Accept': 'application/msword, application/vnd.openxmlformats-officedocument.wordprocessingml.document, */*',
            'Accept-Language': 'de-DE,de;q=0.9,en;q=0.8',
        }

        async with aiohttp.ClientSession(headers=headers) as session:
            # SendenGETKommentar
            logger.info(f"Hinweis{url}")
            async with session.get(url) as response:
                if response.status != 200:
                    logger.error(f"Fehler bei der Verarbeitung{response.status}")
                    return None

                # Kommentar
                content = await response.read()
                content_length = len(content)
                logger.info(f"Hinweisöße: {content_length} Hinweis")

                if content_length == 0:
                    logger.error("Fehler bei der Verarbeitung")
                    return None

                # Kommentar
                with tempfile.NamedTemporaryFile(suffix='.docx', delete=False) as temp_file:
                    temp_file_path = temp_file.name
                    temp_file.write(content)

                logger.info(f"Hinweis{temp_file_path}")

                try:
                    # Kommentar
                    logger.info("Hinweis")
                    doc = Document(temp_file_path)
                    logger.info("DokumenteHinweis")

                    # Kommentar
                    text_content = []
                    for para in doc.paragraphs:
                        if para.text.strip():  # Hinweis
                            text_content.append(para.text)

                    # Kommentar
                    for table in doc.tables:
                        for row in table.rows:
                            for cell in row.cells:
                                if cell.text.strip():  # Hinweis
                                    text_content.append(cell.text)

                    return "\n".join(text_content)
                except Exception as doc_error:
                    logger.debug(f"Fehler bei der Dokumentanalyse: {str(doc_error)}")
                    logger.debug(f"Detaillierte Fehlerinformationen: {traceback.format_exc()}")

                    # Kommentar
                    try:
                        logger.info("Hinweis")
                        # Kommentar
                        import zipfile
                        if zipfile.is_zipfile(temp_file_path):
                            with zipfile.ZipFile(temp_file_path) as zip_ref:
                                # Kommentar
                                logger.info(f"ZIP-Inhalt: {zip_ref.namelist()}")

                                # Kommentar
                                if 'word/document.xml' in zip_ref.namelist():
                                    import xml.etree.ElementTree as ET
                                    with zip_ref.open('word/document.xml') as xml_file:
                                        tree = ET.parse(xml_file)
                                        root = tree.getroot()
                                        # Kommentar
                                        ns = {
                                            'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'
                                        }
                                        text_parts = []
                                        for t in root.findall('.//w:t', ns):
                                            if t.text:
                                                text_parts.append(t.text)
                                        return "\n".join(text_parts)
                        else:
                            logger.debug("Hinweis")
                            return None
                    except Exception as backup_error:
                        logger.error(f"Fehler bei der Verarbeitung{str(backup_error)}")
                        return None
                finally:
                    # Kommentar
                    if temp_file_path and os.path.exists(temp_file_path):
                        os.remove(temp_file_path)
                        logger.info(f"Hinweisöschen: {temp_file_path}")

    except Exception as e:
        logger.error(f"Fehler bei der Word-Verarbeitung: {str(e)}")
        logger.error(f"Detaillierte Fehlerinformationen: {traceback.format_exc()}")
        # Kommentaröschen
        if temp_file_path and os.path.exists(temp_file_path):
            os.remove(temp_file_path)
            logger.info(f"Hinweisöschen: {temp_file_path}")
        return None




async def extract_word_content_from_file(file_path: str) -> Optional[str]:
    """
    Hinweis

    Args:
        file_path: Hinweis

    Returns:
        str: Hinweis
    """
    try:
        logger.info(f"Hinweis{file_path}")
        # Kommentar
        doc = Document(file_path)
        logger.info("DokumenteHinweis")

        # Kommentar
        text_content = []
        for para in doc.paragraphs:
            if para.text.strip():  # Hinweis
                text_content.append(para.text)

        # Kommentar
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    if cell.text.strip():  # Hinweis
                        text_content.append(cell.text)

        return "\n".join(text_content)
    except Exception as doc_error:
        logger.debug(f"Fehler bei der Dokumentanalyse: {str(doc_error)}")
        logger.debug(f"Detaillierte Fehlerinformationen: {traceback.format_exc()}")

        # Kommentar
        try:
            logger.info("Hinweis")
            # Kommentar
            import zipfile
            if zipfile.is_zipfile(file_path):
                with zipfile.ZipFile(file_path) as zip_ref:
                    # Kommentar
                    logger.info(f"ZIP-Inhalt: {zip_ref.namelist()}")

                    # Kommentar
                    if 'word/document.xml' in zip_ref.namelist():
                        import xml.etree.ElementTree as ET
                        with zip_ref.open('word/document.xml') as xml_file:
                            tree = ET.parse(xml_file)
                            root = tree.getroot()
                            # Kommentar
                            ns = {
                                'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'
                            }
                            text_parts = []
                            for t in root.findall('.//w:t', ns):
                                if t.text:
                                    text_parts.append(t.text)
                            return "\n".join(text_parts)
            else:
                logger.debug("Hinweis")
                return None
        except Exception as backup_error:
            logger.error(f"Fehler bei der Verarbeitung{str(backup_error)}")
            return None

    except Exception as e:
        logger.error(f"Fehler bei der Word-Verarbeitung: {str(e)}")
        logger.error(f"Detaillierte Fehlerinformationen: {traceback.format_exc()}")
        return None


async def extract_word_content_from_bytes(file_content: bytes) -> Optional[str]:
    """
    Hinweis

    Args:
        file_content: WordDokumenteHinweis

    Returns:
        str: Hinweis
    """
    
    # Kommentar
    if not file_content or len(file_content) == 0:
        logger.error("Fehler bei der Verarbeitung")
        return "Fehler: Hinweis"
    
    logger.info(f"Hinweisöße: {len(file_content)} Hinweis")
    
    # Kommentar
    try:
        logger.info("Hinweis")
        file_like_object = io.BytesIO(file_content)
        
        doc = Document(file_like_object)
        logger.info("python-docxHinweis")

        # Kommentar
        text_content = []
        
        # Kommentar
        for para in doc.paragraphs:
            if para.text and para.text.strip():
                text_content.append(para.text.strip())

        # Kommentar
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    if cell.text and cell.text.strip():
                        text_content.append(cell.text.strip())
        
        # Kommentar
        for section in doc.sections:
            if section.header:
                for para in section.header.paragraphs:
                    if para.text and para.text.strip():
                        text_content.append(f"[Hinweis{para.text.strip()}")
            if section.footer:
                for para in section.footer.paragraphs:
                    if para.text and para.text.strip():
                        text_content.append(f"[Hinweis{para.text.strip()}")

        if text_content:
            result = "\n".join(text_content)
            logger.info(f"python-docxHinweis{len(result)} Hinweis")
            return result
        else:
            logger.warning("python-docxWarnhinweis")
    
    except Exception as e:
        logger.warning(f"python-docxWarnhinweis{str(e)}")
    
    # Kommentar
    try:
        logger.info("Hinweis")
        import zipfile
        import xml.etree.ElementTree as ET
        
        file_like_object = io.BytesIO(file_content)
        
        if not zipfile.is_zipfile(file_like_object):
            logger.error("Fehler bei der Verarbeitung")
            return "Fehler: Hinweis"
        
        with zipfile.ZipFile(file_like_object) as zf:
            xml_parts = []
            
            # Kommentar
            if 'word/document.xml' in zf.namelist():
                try:
                    with zf.open('word/document.xml') as xml_file:
                        content = xml_file.read()
                        # Kommentar
                        if content.startswith(b'\xef\xbb\xbf'):  # BOM
                            content = content[3:]
                        
                        tree = ET.fromstring(content)
                        
                        # Kommentar
                        namespaces = [
                            {'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'},
                            {'w': 'http://schemas.microsoft.com/office/word/2003/wordml'}
                        ]
                        
                        for ns in namespaces:
                            try:
                                # Kommentar
                                for t in tree.iterfind('.//w:t', ns):
                                    if t.text and t.text.strip():
                                        xml_parts.append(t.text.strip())
                                        
                                # Kommentar
                                for cell in tree.iterfind('.//w:tc', ns):
                                    cell_texts = []
                                    for t in cell.iterfind('.//w:t', ns):
                                        if t.text and t.text.strip():
                                            cell_texts.append(t.text.strip())
                                    if cell_texts:
                                        xml_parts.append(' '.join(cell_texts))
                                        
                                if xml_parts:
                                    break  # Hinweis
                            except Exception:
                                continue
                                
                except ET.ParseError as pe:
                    logger.warning(f"XMLWarnhinweis{pe}")
                except Exception as ee:
                    logger.warning(f"Warnhinweis{ee}")
            
            # Kommentar
            if not xml_parts:
                for filename in zf.namelist():
                    if filename.startswith('word/') and filename.endswith('.xml') and filename != 'word/document.xml':
                        try:
                            with zf.open(filename) as xml_file:
                                tree = ET.parse(xml_file)
                                root = tree.getroot()
                                # Kommentar
                                for elem in root.iter():
                                    if elem.text and elem.text.strip() and len(elem.text.strip()) > 1:
                                        xml_parts.append(elem.text.strip())
                        except Exception:
                            continue
            
            if xml_parts:
                # Kommentar
                unique_parts = list(dict.fromkeys(xml_parts))  # Hinweis
                filtered_parts = [part for part in unique_parts if len(part) > 1]  # Hinweis
                
                if filtered_parts:
                    result = "\n".join(filtered_parts)
                    logger.info(f"XMLHinweis{len(result)} Hinweis")
                    return result
    
    except Exception as e:
        logger.error(f"XMLFehler bei der Verarbeitung{str(e)}")
        logger.error(f"Fehler bei der Verarbeitung{traceback.format_exc()}")
    
    # Kommentar
    if DOCX2TXT_AVAILABLE:
        try:
            logger.info("Hinweis")
            with tempfile.NamedTemporaryFile(suffix='.docx', delete=False) as temp_file:
                temp_file.write(file_content)
                temp_file.flush()
                
                text_content = docx2txt.process(temp_file.name)
                
                # Kommentar
                try:
                    os.unlink(temp_file.name)
                except:
                    pass
                
                if text_content and text_content.strip():
                    result = text_content.strip()
                    logger.info(f"docx2txtHinweis{len(result)} Hinweis")
                    return result
                else:
                    logger.warning("docx2txtWarnhinweis")
        
        except Exception as e:
            logger.warning(f"docx2txtVerarbeitung fehlgeschlagen: {str(e)}")
    
    # Kommentar
    try:
        logger.info("Hinweis")
        file_like_object = io.BytesIO(file_content)
        
        if zipfile.is_zipfile(file_like_object):
            with zipfile.ZipFile(file_like_object) as zf:
                # Kommentar
                text_parts = []
                for filename in zf.namelist():
                    if any(filename.endswith(ext) for ext in ['.xml', '.rels']):
                        try:
                            with zf.open(filename) as f:
                                content = f.read()
                                # Kommentar
                                try:
                                    text_content = content.decode('utf-8')
                                except UnicodeDecodeError:
                                    try:
                                        text_content = content.decode('gbk')
                                    except UnicodeDecodeError:
                                        continue
                                
                                # Kommentar
                                import re
                                # Kommentar
                                text_matches = re.findall(r'>([^<]{2,})<', text_content)
                                for match in text_matches:
                                    cleaned = match.strip()
                                    if len(cleaned) > 2 and not cleaned.isdigit():
                                        text_parts.append(cleaned)
                        except Exception:
                            continue
                
                if text_parts:
                    # Kommentar
                    unique_parts = list(set(text_parts))
                    result = "\n".join(unique_parts)
                    logger.info(f"Hinweis{len(result)} Hinweis")
                    return result
    
    except Exception as e:
        logger.error(f"Fehler bei der Verarbeitung{str(e)}")
    
    # Kommentar
    error_msg = "Fehler bei der Verarbeitung\n1. Fehler bei der Verarbeitung\n2. Fehler bei der Verarbeitung\n3. Fehler bei der Verarbeitung\n4. Fehler bei der Verarbeitung\n\nFehler bei der Verarbeitung"
    logger.error("Fehler bei der Verarbeitung")
    return error_msg
    
# Kommentar
# async def main():
#     url = "https://x/img/2025/05/07/14/6e/6e811b83-8f22-43d2-894b-8fc236ff971f.docx"
#     print(f"Kommentar{url}")
#     content = await extract_word_content_from_url(url)
#     if content:
#         print("Kommentar")
#         print(content)
#     else:
#         print("Kommentar")


# # Kommentar
# if __name__ == "__main__":
#     import asyncio

#     asyncio.run(main())